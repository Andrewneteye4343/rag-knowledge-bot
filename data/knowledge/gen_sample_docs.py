"""產生四種格式的範例知識庫文件（PDF/DOCX/TXT/MD），用於測試不同檔案類型的載入。"""
from pathlib import Path

from docx import Document
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib.units import cm
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.cidfonts import UnicodeCIDFont
from reportlab.platypus import Paragraph, SimpleDocTemplate

OUT = Path("sample_knowledge")
OUT.mkdir(exist_ok=True)

# ---------- 1. PDF：福利制度 ----------
pdf_content = [
    "晨星科技 員工福利制度",
    "本文件說明公司提供的各項員工福利，適用於全體正式員工。",
    "一、團體保險：公司為每位員工投保團體醫療險與意外險，保費全額由公司負擔。",
    "二、年度健康檢查：每年提供一次免費健檢，方案依年齡分級，40 歲以上員工可選擇高階方案。",
    "三、旅遊補助：每年補助員工旅遊經費新台幣 8,000 元，可用於國內外旅遊，需檢附發票核銷。",
    "四、餐飲補助：公司提供午餐補助，每月補助 2,400 元，隨薪資發放。",
    "五、托兒補助：家有 6 歲以下子女之員工，每月可申請托兒補助 3,000 元。",
]
pdfmetrics.registerFont(UnicodeCIDFont("STSong-Light"))
styles = getSampleStyleSheet()
doc = SimpleDocTemplate(str(OUT / "福利制度.pdf"), pagesize=A4)
story = [Paragraph(f"<font name='STSong-Light'>{line}</font>", styles["Normal"]) for line in pdf_content]
doc.build(story)

# ---------- 2. DOCX：教育訓練 ----------
docx = Document()
docx.add_heading("晨星科技 教育訓練制度", level=1)
docx.add_paragraph("本文件說明公司教育訓練的相關規定與補助辦法。")
docx.add_heading("一、新人訓練", level=2)
docx.add_paragraph("新進員工到職後兩週內須完成新人訓練，共 16 小時，內容包含公司簡介、資訊安全與職場倫理。")
docx.add_heading("二、年度訓練時數", level=2)
docx.add_paragraph("每位員工每年至少須完成 24 小時的教育訓練，其中 8 小時須為專業領域課程。")
docx.add_heading("三、外部課程補助", level=2)
docx.add_paragraph("員工參加外部專業課程，公司補助 50% 學費，每人每年上限新台幣 30,000 元。")
docx.add_heading("四、語言學習", level=2)
docx.add_paragraph("公司與線上語言平台合作，員工可免費使用英語學習課程，每月上限 10 小時。")
docx.save(str(OUT / "教育訓練.docx"))

# ---------- 3. TXT：員工活動 ----------
txt_content = """晨星科技 員工活動辦法

本文件說明公司舉辦的各項員工活動與補助規則。

一、家庭日：每年 10 月舉辦家庭日活動，員工可攜帶家屬參加，公司負擔門票與餐飲費用。

二、社團補助：公司鼓勵員工成立社團，每個社團每季可申請補助新台幣 5,000 元，
   需於活動結束後一週內提交成果報告與經費明細。

三、運動會：每年 4 月舉辦員工運動會，設有籃球、羽球與路跑等項目，
   各項比賽前八名可獲得獎金與獎牌。

四、志工活動：公司每年提供 2 天有薪志工假，員工參與公益活動可申請。
"""
(OUT / "員工活動.txt").write_text(txt_content, encoding="utf-8")

# ---------- 4. MD：差旅規定 ----------
md_content = """# 晨星科技 差旅管理辦法

本文件說明員工出差時的費用報支標準與申請流程。

## 一、住宿費標準

- 台北市、新竹市：每晚上限 3,500 元
- 其他縣市：每晚上限 2,800 元
- 國外出差：每晚上限 5,000 元（含早餐）

## 二、交通費

國內出差可搭乘高鐵商務車廂或飛機經濟艙；國外出差機票須提前兩週申請核准。
計程車費用於夜間 10 點後或天候不佳時方可報支。

## 三、出差津貼

國內出差每日伙食津貼 500 元，國外出差每日 1,200 元。
出差期間的網路與國際電話費用，每人每月上限 1,000 元。

## 四、申請流程

出差申請須於出發前一週填寫差旅申請單，經部門主管核准後生效。
費用報支須於出差結束後兩週內完成，檢附發票與登機證。
"""
(OUT / "差旅規定.md").write_text(md_content, encoding="utf-8")

for f in sorted(OUT.iterdir()):
    print(f"  ✓ {f.name}（{f.stat().st_size} bytes）")
print("產生完成！")
